import pandas as pd
import numpy as np
import os

from docx import Document
from docx.shared import Inches
import seaborn as sns
import predictive_plots

def filter_data(df,data_type = np.number,na_number = 50):
    """
    Filters the input DataFrame based on specified conditions: selects columns
    of a specific data type, drops columns with a number of NaN values above a
    certain threshold, and removes the 'Age' column.

    Parameters:
    - df: DataFrame to be filtered.
    - data_type: Data type used to filter the columns. Default is numeric data types.
    - na_number: Threshold for dropping columns based on the count of NaN values.

    Returns:
    - A filtered copy of the input DataFrame.
    """
    # Select columns of the specified data type from the DataFrame
    numeric_columns = df.select_dtypes(include=[data_type]).columns
    columns_to_select = ['EPRIME_CODE'] + list(numeric_columns)

    # Filter out rows where 'clusters' value is negative
    df = df[df["clusters"] >= 0]

    # Select only the relevant columns
    filtered_df = df[columns_to_select]

    # Identify columns to drop based on the threshold of NaN values
    cols_to_drop = filtered_df.columns[filtered_df.isnull().sum() > na_number]
    # Drop identified columns
    filtered_df = filtered_df.drop(columns=cols_to_drop).copy()
    # Drop 'Age' column
    filtered_df.drop("Age", axis=1, inplace=True)

    return filtered_df


def create_metrics_cluster_df(df_cluster, df_scales):
    """
    Merges two DataFrames on 'Subject ID' (for df_cluster) and 'EPRIME_CODE' (for
    df_scales), filters the merged DataFrame, and computes mean values for each
    cluster and metric. It constructs a new DataFrame with these computed means,
    organized by cluster.

    Parameters:
    - df_cluster: DataFrame containing cluster assignments.
    - df_scales: DataFrame containing scale scores.

    Returns:
    - A DataFrame with mean values for each metric, organized by cluster.
    """
    # Merge cluster assignments with scale scores on their common identifiers
    new_data = pd.merge(df_cluster, df_scales, left_on='Subject ID', right_on='EPRIME_CODE')

    # Filter the merged DataFrame
    filter_df = filter_data(new_data)

    # Initialize a dictionary to hold the cluster metrics
    data_clust = {}
    # Iterate over each unique cluster to calculate metrics
    for clust in np.unique(filter_df["clusters"]):
        metric_data = {}
        # Calculate mean for each column (excluding 'EPRIME_CODE' and 'clusters')
        for column in filter_df.drop(["EPRIME_CODE", "clusters"], axis=1).columns:
            mean_col = filter_df[filter_df["clusters"] == clust][column].mean()
            # Store the mean value
            metric_data[column] = {'mean': round(mean_col, 2)}

        # Add the metrics data to the main dictionary
        data_clust[clust] = metric_data

    # Create a DataFrame to hold the organized metrics data
    df_clusters = pd.DataFrame()
    # Populate the DataFrame with the metrics data
    for cluster, attributes in data_clust.items():
        for attribute, values in attributes.items():
            for stat, value in values.items():
                col_name = f"{attribute}_{stat}"
                df_clusters.loc[cluster, col_name] = value

    return df_clusters


def create_mean_tasks(df, df_cluster):
    """
    Merges task scores with cluster assignments, computes mean scores for
    each task category within each cluster, and organizes the results in a
    new DataFrame.

    Parameters:
    - df: DataFrame containing task scores.
    - df_cluster: DataFrame containing cluster assignments.

    Returns:
    - A DataFrame with mean scores for each task category, organized by cluster.
    """
    # Copy the original DataFrame
    copy_data = df.copy()
    # Assign cluster labels to the copied DataFrame
    copy_data["clusters"] = df_cluster["clusters"]

    # Extract unique cluster labels
    clusters = copy_data['clusters'].unique()

    # Initialize an empty DataFrame to store mean scores
    df_mean_scores = pd.DataFrame(columns=['Cluster', 'Happy_0', 'Happy_1', 'Sad_0', 'Sad_1', 'Fear_0', 'Fear_1'])

    # Calculate mean scores for each task category within each cluster
    for cluster in clusters:
        cluster_data = copy_data[copy_data['clusters'] == cluster]
        # Calculate the mean from each individual for each category, and then the mean of all individuals
        # from the same cluster
        df_mean_scores.loc[len(df_mean_scores)] = [
            cluster,
            cluster_data[['Happy_0_' + str(i) for i in range(6)]].mean().mean(),
            cluster_data[['Happy_1_' + str(i) for i in range(9)]].mean().mean(),
            cluster_data[['Sad_0_' + str(i) for i in range(6)]].mean().mean(),
            cluster_data[['Sad_1_' + str(i) for i in range(9)]].mean().mean(),
            cluster_data[['Fear_0_' + str(i) for i in range(6)]].mean().mean(),
            cluster_data[['Fear_1_' + str(i) for i in range(9)]].mean().mean()
        ]

    return df_mean_scores



def create_word(df, list_metrics, doc_name,df_scales = None, cluster_order = None):
    """
    Creates a Word document containing groups of metrics, with each group
    including boxplots and statistical tables for each variable. Plots and tables
    are generated, added to the document, and then the corresponding files are
    deleted.

    Parameters:
    - df: DataFrame containing the data to be visualized and analyzed.
    - list_metrics: List of metric groups, where each group is a list of variable names.
    - doc_name: Name of the output Word document file.

    Note: This function relies on external functions `create_boxplot` and `create_stats_table`
    for generating plots and tables, and requires 'sns' (Seaborn), 'Document' from 'docx',
    and 'Inches' from 'docx.shared' to be imported and configured appropriately.
    """
    sns.set(style="whitegrid")
    # Initialize a new Word document
    doc = Document()
    # Calculate the size of each cluster
    size_clust = np.unique(df["clusters"], return_counts=True)[1]

    # Loop through each group of metrics
    for idx, group in enumerate(list_metrics):
        # Add a paragraph for each group to the document
        doc.add_paragraph(f"Variables {idx + 1}: {', '.join(group)}")

        for i, variable in enumerate(group):
            # Prepare the data for plotting
            col_list = [variable, 'clusters']
            variable_df = df[col_list]
            melted_df = pd.melt(variable_df, id_vars=['clusters'], value_vars=variable,
                                var_name='Variable', value_name='Mean_Score')



            # Generate a boxplot and a statistics table, and save them as images
            plot_filename = predictive_plots.create_boxplot(melted_df, variable, save_plot=True, index=[idx, i],
                                                            cluster_order = cluster_order)
            if df_scales is not None:
                table_filename = predictive_plots.create_stats_table(melted_df, variable, size_clust, save_plot=True,
                                                                     index=[idx, i], df_scales=df_scales,cluster_order = cluster_order)
            else:
                table_filename = predictive_plots.create_stats_table(melted_df, variable, size_clust, save_plot=True,
                                                                     index=[idx, i],cluster_order = cluster_order)


            # Add the images to the Word document
            table_cell = doc.add_table(rows=1, cols=2)
            table_row = table_cell.rows[0]
            cell1, cell2 = table_row.cells
            cell1.paragraphs[0].add_run().add_picture(plot_filename, width=Inches(3))
            cell2.paragraphs[0].add_run().add_picture(table_filename, width=Inches(3))

            # Remove the image files after adding them to the document
            os.remove(plot_filename)
            os.remove(table_filename)

    # Save the Word document
    doc.save(doc_name)


def export_clusters_df(df,output_name):
    """
    Processes the input DataFrame to set specific column names, sorts by 'EPRIME_CODE',
    resets the index, and exports the DataFrame to an Excel file.

    Parameters:
    - df: DataFrame containing 'EPRIME_CODE' and cluster assignments.
    - output_name: Name of the output Excel file.

    Note: This function modifies the column names of the input DataFrame and assumes
    the existence of columns that need to be renamed to 'EPRIME_CODE' and 'clusters'.
    """
    # Rename columns to standard names
    df.columns = ["EPRIME_CODE", "clusters"]
    # Sort DataFrame by 'EPRIME_CODE' and reset index
    df = df.sort_values(by='EPRIME_CODE').reset_index(drop=True)
    # Export the DataFrame to an Excel file
    df.to_excel(output_name, index=False)
