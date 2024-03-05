import pandas as pd
import numpy as np
from scipy.cluster.hierarchy import fcluster

from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import seaborn as sns
import predictive_plots

def filter_data(df,data_type = np.number,na_number = 50):
    numeric_columns = df.select_dtypes(include=[data_type]).columns
    columns_to_select = ['EPRIME_CODE'] + list(numeric_columns)


    # Selecting all numeric columns along with "Subject"
    df = df[df["clusters"] >= 0]

    filtered_df = df[columns_to_select]


    cols_to_drop = filtered_df.columns[filtered_df.isnull().sum() > na_number]
    filtered_df = filtered_df.drop(columns=cols_to_drop).copy()
    filtered_df.drop("Age", axis=1, inplace=True)

    return filtered_df


def create_metrics_cluster_df(df_cluster, df_scales):
    new_data = pd.merge(df_cluster, df_scales, left_on='Subject ID', right_on='EPRIME_CODE')

    filter_df = filter_data(new_data)

    data_clust = {}
    for clust in np.unique(filter_df["clusters"]):
        metric_data = {}
        for column in filter_df.drop(["EPRIME_CODE", "clusters"], axis=1).columns:
            mean_col = filter_df[filter_df["clusters"] == clust][column].mean()
            # std_col = filter_df[filter_df["clusters"] == clust][column].std()
            # max_col = filter_df[filter_df["clusters"] == clust][column].max()
            # min_col = filter_df[filter_df["clusters"] == clust][column].min()

            # Normality Test
            # _, p_value_col = stats.shapiro(result_df[result_df["clusters"]==clust][column])
            metric_data[column] = {'mean': round(mean_col, 2)}  # , 'std': round(std_col, 2),
            # 'max': max_col, 'min': min_col}

        data_clust[clust] = metric_data

    df_clusters = pd.DataFrame()
    for cluster, attributes in data_clust.items():
        for attribute, values in attributes.items():
            for stat, value in values.items():
                col_name = f"{attribute}_{stat}"
                df_clusters.loc[cluster, col_name] = value

    return df_clusters


def create_mean_tasks(df, df_cluster):
    copy_data = df.copy()
    copy_data["clusters"] = df_cluster["clusters"]

    # Extract cluster labels
    clusters = copy_data['clusters'].unique()

    # Create an empty DataFrame to store the mean scores
    df_mean_scores = pd.DataFrame(columns=['Cluster', 'Happy_0', 'Happy_1', 'Sad_0', 'Sad_1', 'Fear_0', 'Fear_1'])

    # Iterate over each cluster
    for cluster in clusters:
        # Filter data for the current cluster
        cluster_data = copy_data[copy_data['clusters'] == cluster]

        # Calculate mean scores for each category
        df_mean_scores.loc[len(df_mean_scores)] = [
            cluster,
            cluster_data[['Happy_0_' + str(i) for i in range(6)]].mean().mean(),
            cluster_data[['Happy_1_' + str(i) for i in range(9)]].mean().mean(),
            cluster_data[['Sad_0_' + str(i) for i in range(6)]].mean().mean(),
            cluster_data[['Sad_1_' + str(i) for i in range(9)]].mean().mean(),
            cluster_data[['Fear_0_' + str(i) for i in range(6)]].mean().mean(),
            cluster_data[['Fear_1_' + str(i) for i in range(9)]].mean().mean()
        ]

    # Display the mean scores for each cluster
    return df_mean_scores



def create_word(df, list_metrics, doc_name):
    sns.set(style="whitegrid")
    doc = Document()
    size_clust = np.unique(df["clusters"], return_counts=True)[1]

    for idx, group in enumerate(list_metrics):
        # Add group name to the document

        doc.add_paragraph(f"Group {idx + 1}: {', '.join(group)}")

        for i, variable in enumerate(group):
            col_list = [variable, 'clusters']
            variable_df = df[col_list]
            melted_df = pd.melt(variable_df, id_vars=['clusters'], value_vars=variable,
                                var_name='Variable', value_name='Mean_Score')

            plot_filename = predictive_plots.create_boxplot(melted_df, variable, save_plot=True, index=[idx, i])

            table_filename = predictive_plots.create_stats_table(melted_df, variable, size_clust, save_plot=True, index=[idx, i])

            # Add plot and table to the Word document in the same row
            table_cell = doc.add_table(rows=1, cols=2)
            table_row = table_cell.rows[0]
            cell1 = table_row.cells[0]
            cell2 = table_row.cells[1]

            cell1.paragraphs[0].add_run().add_picture(plot_filename, width=Inches(3))  # Adjust width as needed
            cell2.paragraphs[0].add_run().add_picture(table_filename, width=Inches(3))  # Adjust width as
    doc.save(doc_name)


def export_clusters_df(df,output_name):

    df.columns = ["EPRIME_CODE", "clusters"]
    df = df.sort_values(by='EPRIME_CODE').reset_index(drop=True)
    df.to_excel(output_name, index=False)
